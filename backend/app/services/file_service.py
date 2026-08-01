"""
FileService
------------
Extracts raw text content from uploaded CV files so the rest of the
pipeline (cv_agent, matcher_agent, etc.) can work with plain text
regardless of whether the user uploaded a .pdf, .docx, or .txt file.
"""

import io
from app.core.logging import logger


class UnsupportedFileTypeError(Exception):
    """Raised when the uploaded file extension/type is not supported."""
    pass


class FileParsingError(Exception):
    """Raised when a file is of a supported type but its text could not be extracted."""
    pass


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


class FileService:
    """Handles extraction of plain text from uploaded resume/CV files."""

    def extract_text(self, filename: str, content: bytes) -> str:
        """
        Extracts text from the given file bytes based on its extension.

        Args:
            filename: The original filename (used to determine file type).
            content: The raw bytes of the uploaded file.

        Returns:
            The extracted, cleaned plain text.

        Raises:
            UnsupportedFileTypeError: If the file extension isn't supported.
            FileParsingError: If the file could not be parsed / contains no text.
        """
        if not filename or "." not in filename:
            raise UnsupportedFileTypeError("File has no extension; cannot determine type.")

        ext = "." + filename.rsplit(".", 1)[-1].lower()

        if ext not in SUPPORTED_EXTENSIONS:
            raise UnsupportedFileTypeError(
                f"Unsupported file type '{ext}'. Supported types: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
            )

        if ext == ".pdf":
            text = self._extract_from_pdf(content)
        elif ext == ".docx":
            text = self._extract_from_docx(content)
        else:  # .txt
            text = self._extract_from_txt(content)

        text = text.strip()
        if not text:
            raise FileParsingError(
                "No extractable text found in the file. "
                "The PDF may be a scanned image without a text layer - "
                "please paste the CV text manually instead."
            )

        return text

    def _extract_from_pdf(self, content: bytes) -> str:
        try:
            from pypdf import PdfReader
        except ImportError as e:
            logger.error("pypdf is not installed; cannot parse PDF uploads.")
            raise FileParsingError("PDF parsing dependency is missing on the server.") from e

        try:
            reader = PdfReader(io.BytesIO(content))
            pages_text = []
            for page in reader.pages:
                page_text = page.extract_text() or ""
                if page_text:
                    pages_text.append(page_text)
            return "\n\n".join(pages_text)
        except Exception as e:
            logger.error(f"Failed to parse PDF: {e}")
            raise FileParsingError(f"Could not read PDF file: {e}") from e

    def _extract_from_docx(self, content: bytes) -> str:
        try:
            import docx
        except ImportError as e:
            logger.error("python-docx is not installed; cannot parse DOCX uploads.")
            raise FileParsingError("DOCX parsing dependency is missing on the server.") from e

        try:
            document = docx.Document(io.BytesIO(content))
            paragraphs = [p.text for p in document.paragraphs if p.text and p.text.strip()]

            # Also pull text out of any tables, since CVs sometimes use table layouts.
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text and cell.text.strip():
                            paragraphs.append(cell.text)

            return "\n".join(paragraphs)
        except Exception as e:
            logger.error(f"Failed to parse DOCX: {e}")
            raise FileParsingError(f"Could not read DOCX file: {e}") from e

    def _extract_from_txt(self, content: bytes) -> str:
        try:
            return content.decode("utf-8")
        except UnicodeDecodeError:
            try:
                return content.decode("latin-1")
            except Exception as e:
                logger.error(f"Failed to decode TXT file: {e}")
                raise FileParsingError(f"Could not read text file: {e}") from e
